from __future__ import annotations

from dataclasses import dataclass
from functools import lru_cache
from pathlib import Path
from uuid import uuid4

from docx import Document as DocxDocument
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader

from app.config import get_settings
from app.llm import get_chat_model, history_to_messages, invoke_with_retry
from app.prompts import build_mode_prompt
from app.schemas import ChatMessage, SourceItem


SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}
_BOOTSTRAP_ATTEMPTED = False


@dataclass
class RetrievedChunk:
    source_id: str
    filename: str
    chunk_index: int
    content: str


@lru_cache
def get_embeddings() -> HuggingFaceEmbeddings:
    settings = get_settings()
    return HuggingFaceEmbeddings(
        model_name=settings.embedding_model_name,
        model_kwargs={"device": settings.embedding_device},
        encode_kwargs={"normalize_embeddings": settings.embedding_normalize},
    )


def get_text_splitter() -> RecursiveCharacterTextSplitter:
    settings = get_settings()
    return RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        separators=["\n\n", "\n", "。", "！", "？", ". ", " ", ""],
    )


def get_vector_store() -> Chroma:
    settings = get_settings()
    settings.vector_store_dir.mkdir(parents=True, exist_ok=True)
    return Chroma(
        collection_name=settings.chroma_collection_name,
        embedding_function=get_embeddings(),
        persist_directory=str(settings.vector_store_dir),
    )


def save_uploaded_file(filename: str, content: bytes) -> Path:
    settings = get_settings()
    settings.raw_data_dir.mkdir(parents=True, exist_ok=True)

    original_name = Path(filename).name or "upload.txt"
    extension = Path(original_name).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise ValueError(f"Unsupported file type: {extension or 'unknown'}. Use {supported}.")

    if not content:
        raise ValueError("Uploaded file is empty.")

    unique_name = f"{uuid4().hex[:8]}_{original_name}"
    target_path = settings.raw_data_dir / unique_name
    target_path.write_bytes(content)
    return target_path


def extract_text(file_path: Path) -> str:
    extension = file_path.suffix.lower()

    if extension == ".txt":
        for encoding in ("utf-8", "utf-8-sig", "gbk"):
            try:
                return file_path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        raise ValueError("Could not decode TXT file with utf-8 or gbk.")

    if extension == ".pdf":
        reader = PdfReader(str(file_path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    if extension == ".docx":
        document = DocxDocument(str(file_path))
        parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        return "\n".join(parts)

    raise ValueError(f"Unsupported file type: {extension or 'unknown'}")


def chunk_text(text: str, chunk_size: int, chunk_overlap: int) -> list[str]:
    cleaned = "\n".join(line.strip() for line in text.splitlines() if line.strip()).strip()
    if not cleaned:
        return []

    if chunk_overlap >= chunk_size:
        raise ValueError("CHUNK_OVERLAP must be smaller than CHUNK_SIZE.")

    splitter = get_text_splitter()
    return [chunk for chunk in splitter.split_text(cleaned) if chunk.strip()]


def ingest_document(file_path: Path, *, vector_store: Chroma | None = None) -> tuple[str, int]:
    settings = get_settings()
    text = extract_text(file_path)
    chunks = chunk_text(text=text, chunk_size=settings.chunk_size, chunk_overlap=settings.chunk_overlap)
    if not chunks:
        raise ValueError("No readable text found in the uploaded document.")

    store = vector_store or get_vector_store()
    document_id = file_path.stem
    ids: list[str] = []
    documents: list[Document] = []

    for index, chunk in enumerate(chunks):
        source_id = f"{document_id}:{index}"
        ids.append(source_id)
        documents.append(
            Document(
                page_content=chunk,
                metadata={
                    "source_id": source_id,
                    "filename": file_path.name,
                    "chunk_index": index,
                },
            )
        )

    try:
        store.delete(ids=ids)
    except Exception:
        pass

    store.add_documents(documents=documents, ids=ids)
    return document_id, len(chunks)


def _ensure_vector_store_bootstrap() -> Chroma:
    global _BOOTSTRAP_ATTEMPTED

    store = get_vector_store()
    if _BOOTSTRAP_ATTEMPTED:
        return store

    _BOOTSTRAP_ATTEMPTED = True
    if store._collection.count() > 0:
        return store

    settings = get_settings()
    if not settings.raw_data_dir.exists():
        return store

    for file_path in settings.raw_data_dir.iterdir():
        if file_path.is_file() and file_path.suffix.lower() in SUPPORTED_EXTENSIONS:
            try:
                ingest_document(file_path, vector_store=store)
            except Exception:
                continue

    return store


def retrieve_chunks(query: str, top_k: int | None = None) -> list[RetrievedChunk]:
    settings = get_settings()
    store = _ensure_vector_store_bootstrap()

    if store._collection.count() == 0:
        return []

    documents = store.similarity_search(query, k=top_k or settings.retrieval_top_k)

    chunks: list[RetrievedChunk] = []
    seen_contents: set[str] = set()
    for document in documents:
        content = document.page_content.strip()
        metadata = document.metadata or {}
        if not content or content in seen_contents:
            continue
        seen_contents.add(content)
        chunks.append(
            RetrievedChunk(
                source_id=str(metadata.get("source_id", "")),
                filename=str(metadata.get("filename", "unknown")),
                chunk_index=int(metadata.get("chunk_index", 0)),
                content=content,
            )
        )

    return chunks


def answer_with_rag(user_message: str, history: list[ChatMessage]) -> tuple[str, str, list[SourceItem]]:
    retrieved_chunks = retrieve_chunks(user_message)
    if not retrieved_chunks:
        raise ValueError("No documents have been ingested yet. Please upload a TXT, PDF, or DOCX file first.")

    context_blocks: list[str] = []
    source_items: list[SourceItem] = []
    for item in retrieved_chunks:
        citation = f"{item.filename}#chunk-{item.chunk_index}"
        context_blocks.append(f"[{citation}]\n{item.content}")
        source_items.append(
            SourceItem(
                source_id=item.source_id,
                filename=item.filename,
                chunk_index=item.chunk_index,
                preview=item.content[:160].replace("\n", " "),
            )
        )

    chain = (
        ChatPromptTemplate.from_messages(
            [
                (
                    "system",
                    (
                        f"{build_mode_prompt('rag')}\n\n"
                        "Answer the user with the retrieved context when relevant. "
                        "If the answer depends on the context, cite sources in square brackets like "
                        "[filename#chunk-0]. If the context is insufficient, say what is missing."
                    ),
                ),
                MessagesPlaceholder("history"),
                (
                    "human",
                    (
                        "Question:\n{message}\n\n"
                        "Retrieved Context:\n{context}\n\n"
                        "Please answer in Chinese when the user uses Chinese."
                    ),
                ),
            ]
        )
        | get_chat_model()
        | StrOutputParser()
    )

    answer = invoke_with_retry(
        "rag_answer_chain",
        lambda: chain.invoke(
            {
                "history": history_to_messages(history),
                "message": user_message,
                "context": "\n\n".join(context_blocks),
            }
        ),
    )
    return answer, get_settings().model_name, source_items
